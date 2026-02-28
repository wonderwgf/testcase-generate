#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
DOCX -> Markdown（用于 Cursor Skills）。
支持自动检测中文路径并使用配置文件方案。
"""

from __future__ import annotations

import argparse
import json
import os
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


# Word XML 命名空间
WORD_NAMESPACE = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': WORD_NAMESPACE}


def _stdout_utf8():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    if hasattr(sys.stderr, "reconfigure"):
        sys.stderr.reconfigure(encoding="utf-8")


_stdout_utf8()


def has_non_ascii(path_str: str) -> bool:
    """检测路径是否包含非 ASCII 字符（如中文）"""
    try:
        path_str.encode('ascii')
        return False
    except UnicodeEncodeError:
        return True


def get_config_file_path() -> Path:
    """获取配置文件的保存路径（用户目录下）"""
    # 优先使用用户目录下的固定位置，避免中文路径
    home = Path.home()
    config_dir = home / ".cursor" / "temp"
    config_dir.mkdir(parents=True, exist_ok=True)
    return config_dir / "docx2md_auto_config.json"


def create_auto_config(docx_path: Path, doc_type: str, out_path: Path = None, images_path: Path = None) -> Path:
    """
    自动创建配置文件以处理中文路径问题。
    返回配置文件路径。
    """
    config = {
        "docx": str(docx_path.resolve()),
        "doc_type": doc_type,
    }
    if out_path:
        config["out"] = str(out_path.resolve())
    if images_path:
        config["images"] = str(images_path.resolve())
    
    config_path = get_config_file_path()
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    print(f"[自动配置] 检测到中文路径，已创建配置文件: {config_path}")
    return config_path


def check_and_handle_chinese_path(args) -> bool:
    """
    检查参数中是否包含中文路径，如果有则自动创建配置文件。
    返回 True 表示已处理（需要从配置文件加载），False 表示无需处理。
    """
    paths_to_check = [
        args.docx,
        args.docx_path_file,
        args.docx_dir,
        args.out,
        args.out_path_file,
        args.images,
    ]
    
    has_chinese = any(has_non_ascii(p) for p in paths_to_check if p)
    
    if not has_chinese:
        return False
    
    print("[自动配置] 检测到参数中包含中文路径，将自动使用配置文件方案...")
    return True


def load_config(config_path: str) -> dict:
    """从配置文件加载参数（UTF-8 编码）"""
    p = Path(config_path)
    if not p.exists():
        raise FileNotFoundError(f"配置文件不存在: {config_path}")
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)


def get_text_from_element(element):
    """
    从 XML 元素中提取文本，包括修订模式（Track Changes）插入的内容。
    处理 <w:ins>（插入）标签中的文本。
    """
    texts = []
    
    # 查找所有文本节点，包括普通文本和修订插入的文本
    # w:t 是普通文本，w:ins/w:r/w:t 是修订插入的文本
    for text_elem in element.iter():
        # 只处理 w:t 标签（文本内容）
        if text_elem.tag == qn('w:t'):
            if text_elem.text:
                texts.append(text_elem.text)
    
    return ''.join(texts)


def get_paragraph_text_with_revisions(paragraph):
    """
    获取段落文本，包含修订模式插入的内容。
    """
    return get_text_from_element(paragraph._element)


def get_runs_with_revisions(paragraph):
    """
    获取段落中的所有 run，包括修订模式（Track Changes）插入的 run。
    返回 (run_element, is_revision) 的列表。
    """
    runs = []
    p_elem = paragraph._element
    
    # 直接遍历段落下的子元素
    for child in p_elem:
        tag = child.tag
        
        # 普通 run
        if tag == qn('w:r'):
            runs.append((child, False))
        
        # 修订模式插入的内容 <w:ins>
        elif tag == qn('w:ins'):
            # ins 标签内部可能包含多个 run
            for ins_child in child:
                if ins_child.tag == qn('w:r'):
                    runs.append((ins_child, True))
    
    return runs


def get_run_text(run_elem):
    """从 run 元素中提取文本"""
    texts = []
    for t_elem in run_elem.findall('.//w:t', NSMAP):
        if t_elem.text:
            texts.append(t_elem.text)
    return ''.join(texts)


def get_run_formatting(run_elem):
    """从 run 元素中提取格式信息"""
    is_bold = False
    is_italic = False
    is_underline = False
    is_strike = False
    color = None
    
    rPr = run_elem.find('w:rPr', NSMAP)
    if rPr is not None:
        # 检查粗体
        bold_elem = rPr.find('w:b', NSMAP)
        if bold_elem is not None:
            val = bold_elem.get(qn('w:val'))
            is_bold = val is None or val.lower() not in ('0', 'false')
        
        # 检查斜体
        italic_elem = rPr.find('w:i', NSMAP)
        if italic_elem is not None:
            val = italic_elem.get(qn('w:val'))
            is_italic = val is None or val.lower() not in ('0', 'false')
        
        # 检查下划线
        underline_elem = rPr.find('w:u', NSMAP)
        if underline_elem is not None:
            val = underline_elem.get(qn('w:val'))
            is_underline = val is not None and val.lower() != 'none'
        
        # 检查删除线
        strike_elem = rPr.find('w:strike', NSMAP)
        if strike_elem is not None:
            val = strike_elem.get(qn('w:val'))
            is_strike = val is None or val.lower() not in ('0', 'false')
        
        # 检查颜色
        color_elem = rPr.find('w:color', NSMAP)
        if color_elem is not None:
            color_val = color_elem.get(qn('w:val'))
            if color_val and color_val.lower() not in ('auto', '000000'):
                color = f"#{color_val}"
                if is_dark_color(color):
                    color = None
    
    return {
        'bold': is_bold,
        'italic': is_italic,
        'underline': is_underline,
        'strike': is_strike,
        'color': color
    }


def format_run_element_text(run_elem):
    """格式化 run 元素的文本（支持修订模式）"""
    text = get_run_text(run_elem)
    if not text:
        return ""
    
    fmt = get_run_formatting(run_elem)
    
    # 修订/批注中的颜色标签不再输出，只保留粗体/斜体/下划线/删除线
    if not any([fmt['bold'], fmt['italic'], fmt['underline'], fmt['strike']]):
        return text
    
    formatted_text = text
    if fmt['strike']:
        formatted_text = f"~~{formatted_text}~~"
    if fmt['underline']:
        formatted_text = f"<u>{formatted_text}</u>"
    if fmt['italic']:
        formatted_text = f"*{formatted_text}*"
    if fmt['bold']:
        formatted_text = f"**{formatted_text}**"
    return formatted_text


def get_paragraph_style_level(paragraph):
    """获取段落的标题层级"""
    style_name = paragraph.style.name
    if "Heading" in style_name or "标题" in style_name:
        if "1" in style_name or "一" in style_name:
            return 1
        if "2" in style_name or "二" in style_name:
            return 2
        if "3" in style_name or "三" in style_name:
            return 3
        if "4" in style_name or "四" in style_name:
            return 4
        if "5" in style_name or "五" in style_name:
            return 5
        if "6" in style_name or "六" in style_name:
            return 6
    return 0


def rgb_to_hex(rgb):
    if rgb is None:
        return None
    try:
        return f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
    except Exception:
        return None


def is_dark_color(color_hex):
    """判断是否是深色/接近黑色的颜色，这些颜色不需要标签（默认颜色）"""
    if not color_hex:
        return True
    hex_color = color_hex.lstrip("#").lower()
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        # RGB 都小于等于 80 视为深色，过滤 #000000、#333333、#262626 等
        if r <= 80 and g <= 80 and b <= 80:
            return True
        return False
    except Exception:
        return False


def get_text_color(run):
    try:
        if run.font.color and run.font.color.rgb:
            color_hex = rgb_to_hex(run.font.color.rgb)
            # 过滤掉深色/接近黑色的颜色（默认颜色不需要标签）
            if is_dark_color(color_hex):
                return None
            return color_hex
    except Exception:
        pass
    return None


def format_run_text(run):
    text = run.text
    if not text:
        return ""

    is_bold = run.bold is True
    is_italic = run.italic is True
    is_underline = run.underline is True
    is_strike = run.font.strike is True

    # 修订/批注中的颜色标签不再输出，只保留粗体/斜体/下划线/删除线
    if not any([is_bold, is_italic, is_underline, is_strike]):
        return text

    formatted_text = text
    if is_strike:
        formatted_text = f"~~{formatted_text}~~"
    if is_underline:
        formatted_text = f"<u>{formatted_text}</u>"
    if is_italic:
        formatted_text = f"*{formatted_text}*"
    if is_bold:
        formatted_text = f"**{formatted_text}**"
    return formatted_text


def extract_images_from_paragraph(paragraph, image_dir, image_counter, output_path=None):
    images_md = []
    try:
        for run in paragraph.runs:
            for drawing in run._element.findall(
                ".//a:blip",
                namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
            ):
                embed = drawing.get(qn("r:embed"))
                if not embed:
                    continue
                try:
                    image_part = paragraph.part.related_parts[embed]
                    image_bytes = image_part.blob

                    content_type = image_part.content_type
                    ext_map = {
                        "image/png": "png",
                        "image/jpeg": "jpg",
                        "image/jpg": "jpg",
                        "image/gif": "gif",
                        "image/bmp": "bmp",
                        "image/tiff": "tiff",
                    }
                    ext = ext_map.get(content_type, "png")

                    image_filename = f"image_{image_counter:03d}.{ext}"
                    image_path = os.path.join(image_dir, image_filename)
                    with open(image_path, "wb") as img_file:
                        img_file.write(image_bytes)

                    if output_path:
                        output_dir = os.path.dirname(os.path.abspath(output_path))
                        try:
                            rel_path = os.path.relpath(image_path, output_dir).replace("\\", "/")
                        except ValueError:
                            rel_path = os.path.join(os.path.basename(image_dir), image_filename).replace("\\", "/")
                    else:
                        rel_path = os.path.join(os.path.basename(image_dir), image_filename).replace("\\", "/")

                    images_md.append(f"![图片{image_counter}]({rel_path})")
                    image_counter += 1
                except Exception as e:
                    print(f"  警告：提取图片失败 - {e}")
    except Exception as e:
        print(f"  警告：处理段落图片时出错 - {e}")
    return images_md, image_counter


def paragraph_to_markdown(paragraph, image_dir, image_counter, output_path=None):
    level = get_paragraph_style_level(paragraph)
    
    # 使用支持修订模式的文本提取方法
    text = get_paragraph_text_with_revisions(paragraph).strip()
    
    # 如果修订模式提取为空，回退到标准方法
    if not text:
        text = paragraph.text.strip()

    images_md, image_counter = extract_images_from_paragraph(paragraph, image_dir, image_counter, output_path)

    if not text and images_md:
        return "\n".join(images_md) + "\n", image_counter
    if not text:
        return "", image_counter

    if level > 0:
        result = f"{'#' * level} {text}\n"
        if images_md:
            result += "\n".join(images_md) + "\n"
        return result, image_counter

    # 获取包含修订的 runs
    runs_with_revisions = get_runs_with_revisions(paragraph)
    
    # 检查是否全部为粗体（用于判断是否为标题）
    if runs_with_revisions:
        all_bold = True
        for run_elem, is_revision in runs_with_revisions:
            run_text = get_run_text(run_elem)
            if run_text.strip():
                fmt = get_run_formatting(run_elem)
                if not fmt['bold']:
                    all_bold = False
                    break
        if all_bold and len(text) < 100:
            result = f"### {text}\n"
            if images_md:
                result += "\n".join(images_md) + "\n"
            return result, image_counter
    elif paragraph.runs:
        # 回退到标准方法
        all_bold = all(run.bold is True for run in paragraph.runs if run.text.strip())
        if all_bold and len(text) < 100:
            result = f"### {text}\n"
            if images_md:
                result += "\n".join(images_md) + "\n"
            return result, image_counter

    if text.startswith(("•", "-", "·")):
        result = f"- {text.lstrip('•-· ')}\n"
        if images_md:
            result += "\n".join(images_md) + "\n"
        return result, image_counter

    if len(text) > 0 and text[0].isdigit() and ("." in text[:5] or "、" in text[:5]):
        result = f"{text}\n"
        if images_md:
            result += "\n".join(images_md) + "\n"
        return result, image_counter

    # 使用支持修订模式的格式化方法
    if runs_with_revisions:
        formatted_text = "".join(format_run_element_text(run_elem) for run_elem, _ in runs_with_revisions).strip()
    else:
        # 回退到标准方法
        formatted_text = "".join(format_run_text(run) for run in paragraph.runs).strip()
    
    result = formatted_text + "\n"
    if images_md:
        result += "\n" + "\n".join(images_md) + "\n"
    return result, image_counter


def format_cell_text(cell):
    """格式化表格单元格文本，支持修订模式"""
    formatted_parts = []
    for paragraph in cell.paragraphs:
        # 优先使用支持修订模式的方法
        runs_with_revisions = get_runs_with_revisions(paragraph)
        if runs_with_revisions:
            for run_elem, _ in runs_with_revisions:
                formatted_parts.append(format_run_element_text(run_elem))
        else:
            # 回退到标准方法
            for run in paragraph.runs:
                formatted_parts.append(format_run_text(run))
    return "".join(formatted_parts).strip().replace("\n", " ")


def convert_table_to_markdown(table):
    if not table.rows:
        return ""
    markdown = "\n"

    header_cells = table.rows[0].cells
    header = "| " + " | ".join(format_cell_text(cell) for cell in header_cells) + " |\n"
    separator = "| " + " | ".join(["---"] * len(header_cells)) + " |\n"
    markdown += header
    markdown += separator

    for row in table.rows[1:]:
        cells = row.cells
        row_text = "| " + " | ".join(format_cell_text(cell) for cell in cells) + " |\n"
        markdown += row_text
    markdown += "\n"
    return markdown


def convert_docx_to_markdown(docx_path, output_path, image_dir):
    """
    将 Word 文档转换为 Markdown，并把图片提取到 image_dir。
    """
    try:
        docx_path_obj = Path(docx_path)
        if not docx_path_obj.is_absolute():
            docx_path = str(docx_path_obj.resolve())
        else:
            docx_path = str(docx_path_obj)

        output_path_obj = Path(output_path)
        if not output_path_obj.is_absolute():
            output_path = str(output_path_obj.resolve())
        else:
            output_path = str(output_path_obj)

        image_dir_obj = Path(image_dir)
        if not image_dir_obj.is_absolute():
            image_dir = str(image_dir_obj.resolve())
        else:
            image_dir = str(image_dir_obj)

        print(f"正在读取Word文档: {docx_path}")
        doc = Document(str(Path(docx_path)))

        os.makedirs(image_dir, exist_ok=True)
        print(f"图片保存目录: {image_dir}")

        markdown_content = []
        image_counter = 1

        for element in doc.element.body:
            if element.tag.endswith("p"):
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        md_text, image_counter = paragraph_to_markdown(
                            paragraph, image_dir, image_counter, output_path
                        )
                        if md_text:
                            markdown_content.append(md_text)
                        break
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        md_table = convert_table_to_markdown(table)
                        if md_table:
                            markdown_content.append(md_table)
                        break

        print(f"正在写入Markdown文件: {output_path}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(markdown_content))

        print("✅ 转换完成！")
        print(f"📄 输出文件: {output_path}")
        print(f"🖼️  提取图片数量: {image_counter - 1}")
        print(f"📊 共转换 {len(markdown_content)} 个元素")
        return True
    except Exception as e:
        print(f"❌ 转换失败: {str(e)}")
        import traceback

        traceback.print_exc()
        return False


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default="", help="配置文件路径（JSON格式，推荐用于中文路径）")
    ap.add_argument("--docx", default="", help="DOCX 路径（包含中文时可改用 --docx-path-file 或 --docx-dir）")
    ap.add_argument("--docx-path-file", default="", help="包含 DOCX 路径的文本文件（UTF-8）")
    ap.add_argument("--docx-dir", default="", help="DOCX 所在目录（避免中文路径编码问题）")
    ap.add_argument("--pattern", default="*.docx", help="从目录中匹配的文件模式（默认 *.docx）")
    ap.add_argument("--docx-index", type=int, default=0, help="从目录匹配结果中选择第 N 个（1 开始）")
    ap.add_argument("--list", action="store_true", help="仅列出目录下匹配文件并退出")
    ap.add_argument("--docx-name-file", default="", help="包含 DOCX 文件名的文本文件（UTF-8，配合 --docx-dir）")
    ap.add_argument("--doc-type", choices=["prd", "design"], default="prd", help="文档类型（影响默认输出目录）")
    ap.add_argument("--out", default="", help="输出 md 路径（可选）")
    ap.add_argument("--out-path-file", default="", help="包含输出 md 路径的文本文件（UTF-8，避免中文路径编码问题）")
    ap.add_argument("--images", default="", help="图片目录（可选）")
    ap.add_argument("--auto-config", action="store_true", help="自动检测中文路径并使用配置文件方案（默认启用）")
    ap.add_argument("--no-auto-config", action="store_true", help="禁用自动配置文件方案")
    args = ap.parse_args()

    # 如果已经指定了 --config，直接从配置文件加载
    if args.config:
        config = load_config(args.config)
        # 仅覆盖配置中显式提供的字段
        for key, value in config.items():
            if not hasattr(args, key):
                continue
            if value is None:
                continue
            if isinstance(value, str) and value == "":
                continue
            setattr(args, key, value)
    # 否则检查是否需要自动创建配置文件（默认启用，除非指定 --no-auto-config）
    elif not args.no_auto_config and check_and_handle_chinese_path(args):
        # 检测到中文路径，需要先解析出 docx_path，然后创建配置文件
        # 这里先尝试解析 docx_path
        docx_path = None
        try:
            if args.docx:
                docx_path = Path(args.docx).expanduser()
                if not docx_path.is_absolute():
                    docx_path = (Path.cwd() / docx_path).resolve()
            elif args.docx_dir:
                docx_dir = Path(args.docx_dir).expanduser()
                if docx_dir.exists():
                    candidates = sorted(docx_dir.glob(args.pattern), key=lambda p: p.stat().st_mtime, reverse=True)
                    if candidates:
                        if args.docx_index > 0 and args.docx_index <= len(candidates):
                            docx_path = candidates[args.docx_index - 1]
                        elif len(candidates) == 1:
                            docx_path = candidates[0]
        except Exception as e:
            print(f"[自动配置] 解析路径时出错: {e}")
        
        if docx_path and docx_path.exists():
            # 推导输出路径
            project_root = None
            for parent in docx_path.parents:
                if parent.name.lower() == "input":
                    project_root = parent.parent
                    break
            if project_root is None:
                project_root = Path.cwd()
            
            out_dir = project_root / "output" / ("prdmd" if args.doc_type == "prd" else "codedesignmd")
            out_dir.mkdir(parents=True, exist_ok=True)
            out_md = out_dir / f"{docx_path.stem}.md"
            img_dir = out_md.parent / f"{out_md.stem}_images"
            
            # 创建配置文件
            config_path = create_auto_config(docx_path, args.doc_type, out_md, img_dir)
            
            # 从配置文件加载参数
            config = load_config(str(config_path))
            for key, value in config.items():
                if not hasattr(args, key):
                    continue
                if value is None:
                    continue
                if isinstance(value, str) and value == "":
                    continue
                setattr(args, key, value)

    docx_path = None
    if args.docx_path_file:
        path_file = Path(args.docx_path_file).expanduser()
        if not path_file.exists():
            raise FileNotFoundError(f"路径文件不存在: {path_file}")
        raw_path = path_file.read_text(encoding="utf-8").strip()
        if not raw_path:
            raise ValueError("路径文件为空")
        docx_path = Path(raw_path).expanduser()
        if not docx_path.is_absolute():
            docx_path = (Path.cwd() / docx_path).resolve()
    elif args.docx:
        docx_path = Path(args.docx).expanduser()
        if not docx_path.is_absolute():
            docx_path = (Path.cwd() / docx_path).resolve()
    elif args.docx_dir:
        docx_dir = Path(args.docx_dir).expanduser()
        if not docx_dir.exists():
            raise FileNotFoundError(f"DOCX 目录不存在: {docx_dir}")
        candidates = sorted(docx_dir.glob(args.pattern), key=lambda p: p.stat().st_mtime, reverse=True)
        if not candidates:
            raise FileNotFoundError(f"目录下未找到匹配文件: {docx_dir} / {args.pattern}")
        if args.list:
            print("匹配到的 DOCX 列表（按修改时间从新到旧）：")
            for i, p in enumerate(candidates, start=1):
                print(f"{i:>2}. {p.name}")
            return 0
        if args.docx_name_file:
            name_file = Path(args.docx_name_file).expanduser()
            if not name_file.exists():
                raise FileNotFoundError(f"文件名文件不存在: {name_file}")
            target_name = name_file.read_text(encoding="utf-8").strip()
            if not target_name:
                raise ValueError("文件名文件为空")
            matched = [p for p in candidates if p.name == target_name]
            if not matched:
                print(f"未找到文件名匹配: {target_name}")
                print("可用文件名：")
                for p in candidates:
                    print(f"- {p.name}")
                return 1
            docx_path = matched[0]
            print(f"按文件名选择: {docx_path}")
        elif args.docx_index > 0:
            if args.docx_index > len(candidates):
                raise IndexError(f"--docx-index 超出范围（1..{len(candidates)}）")
            docx_path = candidates[args.docx_index - 1]
            print(f"按 --docx-index 选择: {docx_path}")
        else:
            if len(candidates) > 1:
                print("目录下有多个 DOCX，请用 --docx-index 选择，或用 --list 查看列表：")
                for i, p in enumerate(candidates, start=1):
                    print(f"{i:>2}. {p.name}")
                return 1
            docx_path = candidates[0]
            print(f"未指定 --docx，自动选择唯一文件: {docx_path}")
    else:
        raise ValueError("请提供 --docx、--docx-path-file 或 --docx-dir")

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX 不存在: {docx_path}")

    # 推导项目根目录：优先根据 .../input/xxx 结构
    project_root = None
    for parent in docx_path.parents:
        if parent.name.lower() == "input":
            project_root = parent.parent
            break
    if project_root is None:
        project_root = Path.cwd()

    if args.out_path_file:
        out_path_file = Path(args.out_path_file).expanduser()
        if not out_path_file.exists():
            raise FileNotFoundError(f"输出路径文件不存在: {out_path_file}")
        raw_out = out_path_file.read_text(encoding="utf-8").strip()
        if not raw_out:
            raise ValueError("输出路径文件为空")
        out_md = Path(raw_out).expanduser()
        if not out_md.is_absolute():
            out_md = (Path.cwd() / out_md).resolve()
    elif args.out:
        out_md = Path(args.out).expanduser().resolve()
    else:
        out_dir = project_root / "output" / ("prdmd" if args.doc_type == "prd" else "codedesignmd")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_md = (out_dir / f"{docx_path.stem}.md").resolve()

    if args.images:
        img_dir = Path(args.images).expanduser().resolve()
    else:
        img_dir = (out_md.parent / f"{out_md.stem}_images").resolve()

    ok = convert_docx_to_markdown(str(docx_path), str(out_md), str(img_dir))
    if not ok:
        raise RuntimeError("转换失败（convert_docx_to_markdown 返回 False）")

    print("转换完成：")
    print(f"- docx: {docx_path}")
    print(f"- md:   {out_md}")
    print(f"- img:  {img_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

